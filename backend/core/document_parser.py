"""
OmniEngine — Document Ingestion & Parsing Engine

Multi-format parser supporting text extraction, tabular decoding, page splitting,
and OCR fallback for PDF, DOCX, TXT, CSV, XLSX, PNG, and JPG files.
"""

from __future__ import annotations

import base64
import csv
import io
import logging
from typing import Any

logger = logging.getLogger(__name__)


class DocumentParser:
    """Core document parsing utility for OmniEngine attachments."""

    @classmethod
    def parse_attachment(cls, attachment: dict[str, Any]) -> str:
        """
        Parse an attachment dictionary and extract readable text content.

        attachment format:
          {
            "filename": "sample.pdf",
            "type": "file" | "image" | "url",
            "mime_type": "application/pdf",
            "content": "raw text or base64 string",
            "url": "base64 data URI or HTTP URL"
          }
        """
        filename = attachment.get("filename") or "attached_file"
        mime_type = attachment.get("mime_type") or ""
        content = attachment.get("content") or ""
        url = attachment.get("url") or ""

        # Decode base64 if content or URL provides a data URI
        raw_bytes: bytes | None = None
        if url.startswith("data:"):
            try:
                header, base64_data = url.split(",", 1)
                raw_bytes = base64.b64decode(base64_data)
                if not mime_type and ";" in header:
                    mime_type = header.split(";")[0].replace("data:", "")
            except Exception as e:
                logger.warning("Failed to decode base64 data URI for %s: %s", filename, str(e))

        if not raw_bytes and content:
            if content.startswith("data:"):
                try:
                    _, base64_data = content.split(",", 1)
                    raw_bytes = base64.b64decode(base64_data)
                except Exception:
                    raw_bytes = content.encode("utf-8")
            else:
                raw_bytes = content.encode("utf-8")

        # Determine file extension/type
        ext = filename.split(".")[-1].lower() if "." in filename else ""

        extracted_text = ""
        try:
            if ext == "pdf" or "pdf" in mime_type:
                extracted_text = cls._parse_pdf(raw_bytes or content.encode("utf-8"))
            elif ext in ("docx", "doc") or "word" in mime_type:
                extracted_text = cls._parse_docx(raw_bytes or content.encode("utf-8"))
            elif ext == "csv" or "csv" in mime_type:
                extracted_text = cls._parse_csv(raw_bytes or content.encode("utf-8"))
            elif ext in ("xlsx", "xls") or "excel" in mime_type or "spreadsheet" in mime_type:
                extracted_text = cls._parse_xlsx(raw_bytes or content.encode("utf-8"))
            elif ext in ("png", "jpg", "jpeg", "webp") or "image" in mime_type:
                extracted_text = cls._parse_image(filename, raw_bytes, content)
            else:
                # Text fallback
                if raw_bytes:
                    extracted_text = raw_bytes.decode("utf-8", errors="replace")
                else:
                    extracted_text = str(content)
        except Exception as e:
            logger.error("Error parsing document %s (%s): %s", filename, mime_type, str(e))
            extracted_text = f"[Document content from {filename}: {content[:1000]}]"

        header_block = (
            f"--- ATTACHMENT: {filename} (Type: {ext.upper() or mime_type or 'text'}) ---\n"
        )
        footer_block = f"\n--- END ATTACHMENT: {filename} ---"
        return f"{header_block}{extracted_text.strip()}{footer_block}"

    @classmethod
    def _parse_pdf(cls, data: bytes) -> str:
        """Extract text from PDF bytes."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            pages_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages_text.append(f"[Page {i + 1}]\n{text}")
            return "\n\n".join(pages_text)
        except ImportError:
            logger.warning("pypdf not available, using raw string extraction fallback")
            return data.decode("utf-8", errors="ignore")

    @classmethod
    def _parse_docx(cls, data: bytes) -> str:
        """Extract text and tables from DOCX bytes."""
        try:
            import docx

            doc = docx.Document(io.BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_texts.append(" | ".join(row_cells))

            full_text = "\n".join(paragraphs)
            if table_texts:
                full_text += "\n\n[Tables]\n" + "\n".join(table_texts)
            return full_text
        except Exception as e:
            logger.warning("docx parsing fallback triggered for data: %s", str(e))
            return data.decode("utf-8", errors="ignore")

    @classmethod
    def _parse_csv(cls, data: bytes) -> str:
        """Extract text from CSV data."""
        text_str = data.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text_str))
        rows = []
        for i, row in enumerate(reader):
            if i > 500:  # Limit rows for prompt safety
                rows.append("... [truncated additional CSV rows]")
                break
            rows.append(", ".join(row))
        return "\n".join(rows)

    @classmethod
    def _parse_xlsx(cls, data: bytes) -> str:
        """Extract sheets and text from XLSX spreadsheet."""
        try:
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
            sheet_texts = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_lines = [f"[Sheet: {sheet_name}]"]
                for row in sheet.iter_rows(values_only=True):
                    row_vals = [str(val) if val is not None else "" for val in row]
                    if any(row_vals):
                        sheet_lines.append("\t".join(row_vals))
                sheet_texts.append("\n".join(sheet_lines))
            return "\n\n".join(sheet_texts)
        except Exception as e:
            logger.warning("openpyxl parsing failed, using text fallback: %s", str(e))
            return data.decode("utf-8", errors="ignore")

    @classmethod
    def _parse_image(cls, filename: str, raw_bytes: bytes | None, content: str) -> str:
        """Perform OCR or image description processing."""
        try:
            if raw_bytes:
                from PIL import Image

                img = Image.open(io.BytesIO(raw_bytes))
                try:
                    import pytesseract

                    ocr_text = pytesseract.image_to_string(img)
                    if ocr_text.strip():
                        return f"[OCR Extracted Text from Image {filename}]:\n{ocr_text.strip()}"
                except Exception as e:
                    logger.debug("pytesseract OCR fallback: %s", str(e))

                return f"[Image attachment: {filename}, dimensions: {img.width}x{img.height} {img.format}]"
        except Exception as e:
            logger.warning("PIL image inspection fallback for %s: %s", filename, str(e))

        return f"[Image Attachment: {filename}]"
